from flask import (
    Blueprint,
    render_template,
    request,
    redirect,
    url_for,
    flash,
    current_app,
    make_response,
)
from flask_login import login_required, current_user
from extensions import db
from models import Atleta, Grupo, AtletaGrupo, AtletaFoto, Responsavel, AtletaResponsavel
from datetime import datetime
import os
import base64
from werkzeug.utils import secure_filename
from urllib.parse import quote

# para exportação em múltiplos formatos
import io
import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from reportlab.lib import colors

atletas_bp = Blueprint("atletas", __name__, url_prefix="/atletas")


def salvar_foto_base64(foto_base64, atleta_id):
    if not foto_base64:
        return

    try:
        header, data = foto_base64.split(",", 1)
    except ValueError:
        return

    try:
        img_bytes = base64.b64decode(data)
    except Exception:
        return

    upload_dir = os.path.join(current_app.root_path, "static", "fotos_atletas")
    os.makedirs(upload_dir, exist_ok=True)

    filename = f"atleta_{atleta_id}_{int(datetime.utcnow().timestamp())}.jpg"
    filepath = os.path.join(upload_dir, secure_filename(filename))

    with open(filepath, "wb") as f:
        f.write(img_bytes)

    rel_path = f"fotos_atletas/{filename}"
    foto = AtletaFoto(atleta_id=atleta_id, foto_path=rel_path)
    db.session.add(foto)


def gerar_arquivo_tabular(nome_base, formato, headers, rows):
    """
    Gera arquivo em CSV, Excel (xlsx), Word (docx) ou PDF a partir
    de headers (lista de strings) e rows (lista de listas).
    """
    formato = (formato or "csv").lower()

    # CSV
    if formato == "csv":
        linhas = [";".join(headers)]
        for r in rows:
            linha = ";".join("" if v is None else str(v) for v in r)
            linhas.append(linha)
        csv_data = "\n".join(linhas)
        resp = make_response(csv_data)
        resp.headers["Content-Type"] = "text/csv; charset=utf-8"
        resp.headers["Content-Disposition"] = f'attachment; filename="{nome_base}.csv"'
        return resp

    # Excel (xlsx)
    if formato == "excel":
        df = pd.DataFrame(rows, columns=headers)
        buf = io.BytesIO()
        with pd.ExcelWriter(buf, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="Dados")
        buf.seek(0)
        resp = make_response(buf.getvalue())
        resp.headers["Content-Type"] = (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        resp.headers["Content-Disposition"] = f'attachment; filename="{nome_base}.xlsx"'
        return resp

    # Word (docx)
    if formato == "word":
        doc = Document()
        doc.add_heading(nome_base, level=1)

        table = doc.add_table(rows=1, cols=len(headers))
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h

        for r in rows:
            row_cells = table.add_row().cells
            for i, v in enumerate(r):
                row_cells[i].text = "" if v is None else str(v)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        resp = make_response(buf.getvalue())
        resp.headers["Content-Type"] = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        resp.headers["Content-Disposition"] = f'attachment; filename="{nome_base}.docx"'
        return resp

    # PDF
    if formato == "pdf":
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4)
        data = [headers] + [[("" if v is None else str(v)) for v in r] for r in rows]
        table = Table(data)
        style = TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e293b")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), 10),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 6),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ]
        )
        table.setStyle(style)
        doc.build([table])
        buf.seek(0)
        resp = make_response(buf.getvalue())
        resp.headers["Content-Type"] = "application/pdf"
        resp.headers["Content-Disposition"] = f'attachment; filename="{nome_base}.pdf"'
        return resp

    # fallback: CSV
    return gerar_arquivo_tabular(nome_base, "csv", headers, rows)


@atletas_bp.route("/listar")
@login_required
def listar():
    if current_user.role == "PARENT":
        links = (
            AtletaResponsavel.query.join(AtletaResponsavel.responsavel)
            .filter(Responsavel.usuario_id == current_user.id)
            .all()
        )
        ids = [l.atleta_id for l in links]
        if ids:
            atletas = Atleta.query.filter(Atleta.id.in_(ids)).order_by(Atleta.nome).all()
        else:
            atletas = []
    else:
        atletas = Atleta.query.order_by(Atleta.nome).all()

    return render_template("atletas_listar.html", atletas=atletas, filtros={})


@atletas_bp.route("/novo", methods=["GET", "POST"])
@login_required
def novo():
    grupos = Grupo.query.order_by(Grupo.nome).all()

    if request.method == "POST":
        nome = (request.form.get("nome") or "").strip()
        dn = request.form.get("data_nascimento")
        posicao = (request.form.get("posicao") or "").strip() or None

        rg = (request.form.get("rg") or "").strip() or None
        cpf = (request.form.get("cpf") or "").strip() or None

        telefone = (request.form.get("telefone") or "").strip() or None
        validade_atestado_str = request.form.get("validade_atestado") or None
        info_adicionais = (request.form.get("informacoes_adicionais") or "").strip() or None

        resp_nome = (request.form.get("responsavel_nome") or "").strip() or None
        resp_cpf = (request.form.get("responsavel_cpf") or "").strip() or None
        resp_tel = (request.form.get("responsavel_telefone") or "").strip() or None
        resp_parentesco = (request.form.get("responsavel_parentesco") or "").strip() or None

        grupos_ids = request.form.getlist("grupos_ids")
        foto_base64 = request.form.get("foto_base64")

        if not nome or not dn:
            flash("Preencha nome e data de nascimento.", "danger")
            return redirect(url_for("atletas.novo"))

        try:
            data_nasc = datetime.strptime(dn, "%Y-%m-%d").date()
        except ValueError:
            flash("Data de nascimento inválida.", "danger")
            return redirect(url_for("atletas.novo"))

        validade_atestado = None
        if validade_atestado_str:
            try:
                validade_atestado = datetime.strptime(validade_atestado_str, "%Y-%m-%d").date()
            except ValueError:
                flash("Data de validade do atestado inválida.", "danger")
                return redirect(url_for("atletas.novo"))

        atleta = Atleta(
            nome=nome,
            data_nascimento=data_nasc,
            posicao=posicao,
            rg=rg,
            cpf=cpf,
            telefone=telefone,
            validade_atestado=validade_atestado,
            informacoes_adicionais=info_adicionais,
            responsavel_nome=resp_nome,
            responsavel_cpf=resp_cpf,
            responsavel_telefone=resp_tel,
            responsavel_parentesco=resp_parentesco,
            status="ATIVO",
        )

        db.session.add(atleta)
        db.session.flush()

        for gid in grupos_ids:
            if gid:
                ag = AtletaGrupo(atleta_id=atleta.id, grupo_id=int(gid), ativo=True)
                db.session.add(ag)

        if foto_base64:
            salvar_foto_base64(foto_base64, atleta.id)

        db.session.commit()
        flash("Atleta cadastrado com sucesso!", "success")
        return redirect(url_for("atletas.listar"))

    return render_template("atletas_novo.html", atleta=None, grupos=grupos, grupos_do_atleta=set(), foto_atual=None)


@atletas_bp.route("/<int:atleta_id>/editar", methods=["GET", "POST"])
@login_required
def editar(atleta_id):
    atleta = Atleta.query.get_or_404(atleta_id)
    grupos = Grupo.query.order_by(Grupo.nome).all()
    grupos_do_atleta = {ag.grupo_id for ag in atleta.grupos}

    if request.method == "POST":
        nome = (request.form.get("nome") or "").strip()
        dn = request.form.get("data_nascimento")
        posicao = (request.form.get("posicao") or "").strip() or None

        rg = (request.form.get("rg") or "").strip() or None
        cpf = (request.form.get("cpf") or "").strip() or None

        telefone = (request.form.get("telefone") or "").strip() or None
        validade_atestado_str = request.form.get("validade_atestado") or None
        info_adicionais = (request.form.get("informacoes_adicionais") or "").strip() or None

        resp_nome = (request.form.get("responsavel_nome") or "").strip() or None
        resp_cpf = (request.form.get("responsavel_cpf") or "").strip() or None
        resp_tel = (request.form.get("responsavel_telefone") or "").strip() or None
        resp_parentesco = (request.form.get("responsavel_parentesco") or "").strip() or None

        grupos_ids = request.form.getlist("grupos_ids")
        foto_base64 = request.form.get("foto_base64")
        status = request.form.get("status") or atleta.status

        if not nome or not dn:
            flash("Preencha nome e data de nascimento.", "danger")
            return redirect(url_for("atletas.editar", atleta_id=atleta.id))

        try:
            data_nasc = datetime.strptime(dn, "%Y-%m-%d").date()
        except ValueError:
            flash("Data de nascimento inválida.", "danger")
            return redirect(url_for("atletas.editar", atleta_id=atleta.id))

        validade_atestado = None
        if validade_atestado_str:
            try:
                validade_atestado = datetime.strptime(validade_atestado_str, "%Y-%m-%d").date()
            except ValueError:
                flash("Data de validade do atestado inválida.", "danger")
                return redirect(url_for("atletas.editar", atleta_id=atleta.id))

        atleta.nome = nome
        atleta.data_nascimento = data_nasc
        atleta.posicao = posicao
        atleta.rg = rg
        atleta.cpf = cpf
        atleta.telefone = telefone
        atleta.validade_atestado = validade_atestado
        atleta.informacoes_adicionais = info_adicionais
        atleta.responsavel_nome = resp_nome
        atleta.responsavel_cpf = resp_cpf
        atleta.responsavel_telefone = resp_tel
        atleta.responsavel_parentesco = resp_parentesco
        atleta.status = status

        AtletaGrupo.query.filter_by(atleta_id=atleta.id).delete()
        for gid in grupos_ids:
            if gid:
                ag = AtletaGrupo(atleta_id=atleta.id, grupo_id=int(gid), ativo=True)
                db.session.add(ag)

        if foto_base64:
            salvar_foto_base64(foto_base64, atleta.id)

        db.session.commit()
        flash("Dados do atleta atualizados!", "success")
        return redirect(url_for("atletas.listar"))

    foto_atual = atleta.fotos[-1].foto_path if atleta.fotos else None

    return render_template(
        "atletas_novo.html",
        atleta=atleta,
        grupos=grupos,
        grupos_do_atleta=grupos_do_atleta,
        foto_atual=foto_atual,
    )


@atletas_bp.route("/exportar")
@login_required
def exportar():
    formato = request.args.get("formato", "csv")

    atletas = Atleta.query.order_by(Atleta.nome).all()
    headers = ["Nome", "Data de nascimento", "Posição", "Responsável", "Telefone"]

    rows = []
    for a in atletas:
        dn = a.data_nascimento.strftime("%d/%m/%Y") if a.data_nascimento else ""
        telefone = a.responsavel_telefone or a.telefone or ""
        rows.append(
            [
                a.nome or "",
                dn,
                a.posicao or "",
                a.responsavel_nome or "",
                telefone,
            ]
        )

    return gerar_arquivo_tabular("alunos", formato, headers, rows)


@atletas_bp.route("/importar", methods=["GET", "POST"])
@login_required
def importar():
    if request.method == "POST":
        flash("Importação em desenvolvimento. Em breve você poderá subir uma planilha com os alunos.", "info")
        return redirect(url_for("atletas.importar"))

    return render_template("atletas_importar.html")


@atletas_bp.route("/cobrar-documentos/<int:atleta_id>")
@login_required
def cobrar_documentos(atleta_id):
    atleta = Atleta.query.get_or_404(atleta_id)

    nome_resp = atleta.responsavel_nome or "responsável"
    nome_atl = atleta.nome

    docs = [
        "atestado de matrícula escolar",
        "comprovante de residência",
        "cópia do RG",
        "atestado médico (validade de 1 ano)",
    ]

    msg = (
        f"Olá {nome_resp}, tudo bem?\n\n"
        f"O seu filho(a) {nome_atl} está com os seguintes documentos pendentes:\n"
        + "\n".join(f"- {d}" for d in docs)
        + f"\n\nEnvie uma foto de cada documento ou leve tudo no próximo treino do(a) {nome_atl}."
    )

    telefone = (atleta.responsavel_telefone or atleta.telefone or "").strip()
    telefone_digits = "".join(ch for ch in telefone if ch.isdigit())

    if not telefone_digits:
        flash("Não há telefone cadastrado para este responsável.", "danger")
        return redirect(url_for("atletas.listar"))

    url = f"https://wa.me/55{telefone_digits}?text={quote(msg)}"
    return redirect(url)
