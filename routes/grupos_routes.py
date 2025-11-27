from flask import Blueprint, render_template, request, redirect, url_for, flash, make_response
from flask_login import login_required, current_user
from extensions import db
from models import Grupo, Atleta, AtletaGrupo

import io
import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from reportlab.lib import colors


grupos_bp = Blueprint("grupos", __name__, url_prefix="/grupos")


def _require_staff():
    return current_user.role in ("ADMIN", "COACH", "SUPER_ADMIN")


def gerar_arquivo_tabular(nome_base, formato, headers, rows):
    formato = (formato or "csv").lower()

    if formato == "csv":
        linhas = [";".join(headers)]
        for r in rows:
            linha = ";".join("" if v is None else str(v) for v in r)
            linhas.append(linha)
        csv_data = "\n".join(linhas)
        resp = make_response(csv_data)
        resp.headers["Content-Type"] = "text/csv; charset=utf-8"
        resp.headers["Content-Disposition"] = f'attachment; filename="{nome_base}.csv"'
        return resp

    if formato == "excel":
        df = pd.DataFrame(rows, columns=headers)
        buf = io.BytesIO()
        with pd.ExcelWriter(buf, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="Dados")
        buf.seek(0)
        resp = make_response(buf.getvalue())
        resp.headers["Content-Type"] = (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        resp.headers["Content-Disposition"] = f'attachment; filename="{nome_base}.xlsx"'
        return resp

    if formato == "word":
        doc = Document()
        doc.add_heading(nome_base, level=1)
        table = doc.add_table(rows=1, cols=len(headers))
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
        for r in rows:
            row_cells = table.add_row().cells
            for i, v in enumerate(r):
                row_cells[i].text = "" if v is None else str(v)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        resp = make_response(buf.getvalue())
        resp.headers["Content-Type"] = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        resp.headers["Content-Disposition"] = f'attachment; filename="{nome_base}.docx"'
        return resp

    if formato == "pdf":
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4)
        data = [headers] + [[("" if v is None else str(v)) for v in r] for r in rows]
        table = Table(data)
        style = TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e293b")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), 10),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 6),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ]
        )
        table.setStyle(style)
        doc.build([table])
        buf.seek(0)
        resp = make_response(buf.getvalue())
        resp.headers["Content-Type"] = "application/pdf"
        resp.headers["Content-Disposition"] = f'attachment; filename="{nome_base}.pdf"'
        return resp

    return gerar_arquivo_tabular(nome_base, "csv", headers, rows)


@grupos_bp.route("/listar")
@login_required
def listar():
    if not _require_staff():
        flash("Você não tem permissão para acessar os grupos.", "danger")
        return redirect(url_for("dashboard.index"))

    grupos = Grupo.query.order_by(Grupo.nome).all()
    return render_template("grupos_listar.html", grupos=grupos)


@grupos_bp.route("/novo", methods=["GET", "POST"])
@login_required
def novo():
    if not _require_staff():
        flash("Você não tem permissão para criar grupos.", "danger")
        return redirect(url_for("dashboard.index"))

    if request.method == "POST":
        nome = (request.form.get("nome") or "").strip()
        faixa_min = request.form.get("faixa_etaria_min") or None
        faixa_max = request.form.get("faixa_etaria_max") or None
        descricao = (request.form.get("descricao") or "").strip()

        if not nome:
            flash("Informe o nome do grupo.", "danger")
            return redirect(url_for("grupos.novo"))

        grupo = Grupo(
            nome=nome,
            faixa_etaria_min=int(faixa_min) if faixa_min else None,
            faixa_etaria_max=int(faixa_max) if faixa_max else None,
            descricao=descricao,
        )
        db.session.add(grupo)
        db.session.commit()

        flash("Grupo criado com sucesso.", "success")
        return redirect(url_for("grupos.listar"))

    return render_template("grupos_form.html", grupo=None)


@grupos_bp.route("/<int:grupo_id>/editar", methods=["GET", "POST"])
@login_required
def editar(grupo_id):
    if not _require_staff():
        flash("Você não tem permissão para editar grupos.", "danger")
        return redirect(url_for("dashboard.index"))

    grupo = Grupo.query.get_or_404(grupo_id)

    if request.method == "POST":
        nome = (request.form.get("nome") or "").strip()
        faixa_min = request.form.get("faixa_etaria_min") or None
        faixa_max = request.form.get("faixa_etaria_max") or None
        descricao = (request.form.get("descricao") or "").strip()

        if not nome:
            flash("Informe o nome do grupo.", "danger")
            return redirect(url_for("grupos.editar", grupo_id=grupo.id))

        grupo.nome = nome
        grupo.faixa_etaria_min = int(faixa_min) if faixa_min else None
        grupo.faixa_etaria_max = int(faixa_max) if faixa_max else None
        grupo.descricao = descricao

        db.session.commit()
        flash("Grupo atualizado com sucesso.", "success")
        return redirect(url_for("grupos.listar"))

    return render_template("grupos_form.html", grupo=grupo)


@grupos_bp.route("/exportar")
@login_required
def exportar():
    if not _require_staff():
        flash("Você não tem permissão para exportar grupos.", "danger")
        return redirect(url_for("dashboard.index"))

    formato = request.args.get("formato", "csv")

    grupos = Grupo.query.order_by(Grupo.nome).all()
    headers = ["Nome", "Faixa etária mínima", "Faixa etária máxima", "Descrição"]
    rows = []
    for g in grupos:
        rows.append(
            [
                g.nome or "",
                g.faixa_etaria_min or "",
                g.faixa_etaria_max or "",
                (g.descricao or "").replace("\n", " "),
            ]
        )

    return gerar_arquivo_tabular("grupos", formato, headers, rows)


@grupos_bp.route("/<int:grupo_id>/exportar")
@login_required
def exportar_grupo(grupo_id):
    """
    Exporta um único grupo com todos os atletas vinculados
    (nome, posição, CPF/RG, data de nascimento, responsável, telefone).
    """
    if not _require_staff():
        flash("Você não tem permissão para exportar grupos.", "danger")
        return redirect(url_for("dashboard.index"))

    formato = request.args.get("formato", "csv")
    grupo = Grupo.query.get_or_404(grupo_id)

    # Busca atletas vinculados ao grupo
    atletas = (
        Atleta.query.join(AtletaGrupo, Atleta.id == AtletaGrupo.atleta_id)
        .filter(AtletaGrupo.grupo_id == grupo.id)
        .order_by(Atleta.nome)
        .all()
    )

    headers = [
        "Grupo",
        "Nome do atleta",
        "Posição",
        "CPF",
        "RG",
        "Data de nascimento",
        "Responsável",
        "Telefone",
    ]
    rows = []
    for a in atletas:
        dn = a.data_nascimento.strftime("%d/%m/%Y") if getattr(a, "data_nascimento", None) else ""
        telefone = a.responsavel_telefone or a.telefone or ""
        rows.append(
            [
                grupo.nome,
                a.nome or "",
                a.posicao or "",
                a.cpf or "",
                a.rg or "",
                dn,
                a.responsavel_nome or "",
                telefone,
            ]
        )

    nome_base = f"grupo_{grupo.id}_{grupo.nome.replace(' ', '_')}"
    return gerar_arquivo_tabular(nome_base, formato, headers, rows)
