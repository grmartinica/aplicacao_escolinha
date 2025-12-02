from flask import (
    Blueprint,
    render_template,
    request,
    redirect,
    url_for,
    flash,
    make_response,
)
from flask_login import login_required, current_user

from extensions import db
from models import (
    ContaReceber,
    AtletaResponsavel,
    Atleta,
    Responsavel,
    ContaPagar,
    AtletaPlano,
    Plano,
)
from urllib.parse import quote
from datetime import datetime, date, timedelta
import os

# exportação multi-formato
import io
import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from reportlab.lib import colors


financeiro_bp = Blueprint("financeiro", __name__, url_prefix="/financeiro")


def _require_admin():
    if current_user.role not in ("ADMIN", "SUPER_ADMIN"):
        flash("Você não tem permissão para essa operação financeira.", "danger")
        return False
    return True


def gerar_arquivo_tabular(nome_base, formato, headers, rows):
    formato = (formato or "csv").lower()

    # CSV
    if formato == "csv":
        linhas = [";".join(headers)]
        for r in rows:
            linha = ";".join("" if v is None else str(v) for v in r)
            linhas.append(linha)
        csv_data = "\n".join(linhas)
        resp = make_response(csv_data)
        resp.headers["Content-Type"] = "text/csv; charset=utf-8"
        resp.headers["Content-Disposition"] = f'attachment; filename="{nome_base}.csv"'
        return resp

    # Excel
    if formato == "excel":
        df = pd.DataFrame(rows, columns=headers)
        buf = io.BytesIO()
        with pd.ExcelWriter(buf, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="Dados")
        buf.seek(0)
        resp = make_response(buf.getvalue())
        resp.headers["Content-Type"] = (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        resp.headers["Content-Disposition"] = f'attachment; filename="{nome_base}.xlsx"'
        return resp

    # Word
    if formato == "word":
        doc = Document()
        doc.add_heading(nome_base, level=1)
        table = doc.add_table(rows=1, cols=len(headers))
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
        for r in rows:
            row_cells = table.add_row().cells
            for i, v in enumerate(r):
                row_cells[i].text = "" if v is None else str(v)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        resp = make_response(buf.getvalue())
        resp.headers["Content-Type"] = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        resp.headers["Content-Disposition"] = f'attachment; filename="{nome_base}.docx"'
        return resp

    # PDF
    if formato == "pdf":
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4)
        data = [headers] + [[("" if v is None else str(v)) for v in r] for r in rows]
        table = Table(data)
        style = TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e293b")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), 10),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 6),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ]
        )
        table.setStyle(style)
        doc.build([table])
        buf.seek(0)
        resp = make_response(buf.getvalue())
        resp.headers["Content-Type"] = "application/pdf"
        resp.headers["Content-Disposition"] = f'attachment; filename="{nome_base}.pdf"'
        return resp

    # fallback
    return gerar_arquivo_tabular(nome_base, "csv", headers, rows)


# =========================
# Cobrança automática
# =========================

def primeiro_dia_util(ano: int, mes: int) -> date:
    """
    Retorna o primeiro dia útil (segunda a sexta) do mês/ano informado.
    """
    d = date(ano, mes, 1)
    # weekday(): 0=segunda ... 6=domingo
    while d.weekday() >= 5:  # 5=sábado, 6=domingo
        d += timedelta(days=1)
    return d


def _gerar_cobrancas_para_mes(ano: int, mes: int) -> int:
    """
    Gera cobranças (ContaReceber) para todos atletas com plano ativo
    na competência ano/mes. Não duplica se já existir para o mês.
    Retorna a quantidade criada.
    """
    competencia = date(ano, mes, 1)
    planos_ativos = AtletaPlano.query.filter_by(ativo=True).all()
    criadas = 0

    for ap in planos_ativos:
        plano = ap.plano
        atleta = ap.atleta

        if not plano or not atleta:
            continue

        # não gera duplicado
        existente = ContaReceber.query.filter_by(
            atleta_id=atleta.id,
            competencia=competencia,
        ).first()
        if existente:
            continue

        dia_venc = plano.dia_vencimento or 10
        vencimento = date(ano, mes, dia_venc)

        conta = ContaReceber(
            atleta_id=atleta.id,
            descricao=f"Mensalidade {plano.nome} {competencia.strftime('%m/%Y')}",
            competencia=competencia,
            vencimento=vencimento,
            valor=plano.valor_mensal,
            status="PENDENTE",
            metodo_pagamento=plano.forma_pagamento_padrao,
        )
        db.session.add(conta)
        criadas += 1

    if criadas > 0:
        db.session.commit()
    return criadas


# guarda em memória qual competência já processamos
_ultima_competencia_processada = None


@financeiro_bp.before_app_request
def _auto_gerar_cobrancas_mes_atual():
    """
    Roda automaticamente a geração de cobranças no primeiro dia útil
    do mês, na primeira requisição autenticada que bater na aplicação.

    - Só roda se:
        * usuário estiver logado (qualquer papel)
        * hoje >= primeiro dia útil do mês
        * ainda não processamos esta competência neste processo
    """
    global _ultima_competencia_processada

    if not current_user.is_authenticated:
        return

    hoje = date.today()
    competencia = date(hoje.year, hoje.month, 1)
    primeiro_util = primeiro_dia_util(hoje.year, hoje.month)

    # ainda não chegou o primeiro dia útil -> não gera nada
    if hoje < primeiro_util:
        return

    if _ultima_competencia_processada == competencia:
        return

    criadas = _gerar_cobrancas_para_mes(hoje.year, hoje.month)
    _ultima_competencia_processada = competencia

    # sem flash aqui pra não poluir a experiência do usuário final
    # (se quiser logar, pode usar print ou logging)


@financeiro_bp.route("/gerar_cobrancas_automaticas")
@login_required
def gerar_cobrancas_automaticas():
    """
    Rota manual (fallback). Se quiser, nem precisa mostrar botão na interface.
    """
    if not _require_admin():
        return redirect(url_for("financeiro.resumo"))

    hoje = date.today()
    competencia = date(hoje.year, hoje.month, 1)

    criadas = _gerar_cobrancas_para_mes(hoje.year, hoje.month)

    if criadas == 0:
        flash(
            f"Nenhuma cobrança nova criada para {competencia.strftime('%m/%Y')} "
            f"(já existiam registros).",
            "info",
        )
    else:
        flash(
            f"Cobranças geradas para {criadas} atleta(s) em {competencia.strftime('%m/%Y')}.",
            "success",
        )

    return redirect(url_for("financeiro.resumo"))


# =========================
# RESUMO FINANCEIRO
# =========================

@financeiro_bp.route("/resumo")
@login_required
def resumo():
    # Se não for admin / super_admin, mostra só a visão "leiaute" sem ações
    if current_user.role not in ("ADMIN", "SUPER_ADMIN"):
        return render_template(
            "financeiro_resumo.html",
            total_pendente=0,
            total_pago=0,
            qtd_inad=0,
            itens=[],
            despesas=[],
            filtros={"nome": "", "status": "TODOS", "inadimplentes": "0"},
        )

    nome = (request.args.get("nome") or "").strip()
    status = request.args.get("status") or "TODOS"
    so_inadimplentes = request.args.get("inadimplentes") == "1"

    query = ContaReceber.query.join(ContaReceber.atleta, isouter=True)

    if so_inadimplentes:
        query = query.filter(ContaReceber.status.in_(("PENDENTE", "ATRASADO")))
    elif status != "TODOS":
        query = query.filter(ContaReceber.status == status)

    if nome:
        query = query.filter(Atleta.nome.ilike(f"%{nome}%"))

    itens = query.order_by(ContaReceber.vencimento.desc()).limit(200).all()

    total_pendente = (
        db.session.query(db.func.sum(ContaReceber.valor))
        .filter(ContaReceber.status != "PAGO")
        .scalar()
        or 0
    )
    total_pago = (
        db.session.query(db.func.sum(ContaReceber.valor))
        .filter(ContaReceber.status == "PAGO")
        .scalar()
        or 0
    )
    qtd_inad = (
        ContaReceber.query.filter(
            ContaReceber.status.in_(("PENDENTE", "ATRASADO"))
        ).count()
    )

    despesas = ContaPagar.query.order_by(ContaPagar.vencimento.desc()).limit(200).all()

    filtros = {
        "nome": nome,
        "status": status,
        "inadimplentes": "1" if so_inadimplentes else "0",
    }

    return render_template(
        "financeiro_resumo.html",
        total_pendente=float(total_pendente),
        total_pago=float(total_pago),
        qtd_inad=qtd_inad,
        itens=itens,
        despesas=despesas,
        filtros=filtros,
    )


@financeiro_bp.route("/responsavel")
@login_required
def resumo_responsavel():
    if current_user.role != "PARENT":
        return render_template("financeiro_responsavel.html", itens=[])

    links = (
        AtletaResponsavel.query.join(AtletaResponsavel.responsavel)
        .filter(AtletaResponsavel.responsavel.has(usuario_id=current_user.id))
        .all()
    )
    atleta_ids = [l.atleta_id for l in links]
    itens = (
        ContaReceber.query.filter(ContaReceber.atleta_id.in_(atleta_ids))
        .order_by(ContaReceber.vencimento.desc())
        .all()
    )
    return render_template("financeiro_responsavel.html", itens=itens)


@financeiro_bp.route("/exportar")
@login_required
def exportar():
    if not _require_admin():
        return redirect(url_for("financeiro.resumo"))

    formato = request.args.get("formato", "csv")

    itens = (
        ContaReceber.query.join(ContaReceber.atleta, isouter=True)
        .order_by(ContaReceber.vencimento.desc())
        .all()
    )

    headers = ["Aluno", "Vencimento", "Valor", "Status"]
    rows = []
    for c in itens:
        nome = c.atleta.nome if c.atleta else ""
        ven = c.vencimento.strftime("%d/%m/%Y") if c.vencimento else ""
        rows.append(
            [
                nome,
                ven,
                f"{float(c.valor or 0):.2f}",
                c.status or "",
            ]
        )

    return gerar_arquivo_tabular("financeiro", formato, headers, rows)


# =========================
# COBRANÇA VIA WHATSAPP
# =========================

@financeiro_bp.route("/cobrar/<int:atleta_id>")
@login_required
def cobrar_whatsapp(atleta_id):
    """
    Cobrança via WhatsApp:
    - Soma pendências do atleta
    - Cria um pagamento Pix único no Mercado Pago com valor total (se configurado)
    - Inclui o link dessa cobrança na mensagem do WhatsApp
    - Ajuste para funcionar mesmo sem contato salvo (normalização do telefone)
    """
    if not _require_admin():
        return redirect(url_for("financeiro.resumo"))

    atleta = Atleta.query.get_or_404(atleta_id)

    pendencias = (
        ContaReceber.query.filter(
            ContaReceber.atleta_id == atleta_id,
            ContaReceber.status.in_(("PENDENTE", "ATRASADO")),
        )
        .order_by(ContaReceber.vencimento)
        .all()
    )

    if not pendencias:
        flash("Este atleta não possui pendências financeiras.", "info")
        return redirect(url_for("financeiro.resumo"))

    meses = []
    total = 0.0
    for c in pendencias:
        if c.vencimento:
            meses.append(c.vencimento.strftime("%m/%Y"))
        total += float(c.valor or 0)

    meses_str = ", ".join(sorted(set(meses)))
    nome_resp = atleta.responsavel_nome or "responsável"
    nome_atl = atleta.nome

    # Gera um pagamento Pix único via Mercado Pago (se configurado)
    link_pagamento = ""
    try:
        import mercadopago

        access_token = os.getenv("MERCADOPAGO_ACCESS_TOKEN")
        if access_token:
            sdk = mercadopago.SDK(access_token)
            description = f"Mensalidades em aberto - {nome_atl} ({meses_str})"

            payment_data = {
                "transaction_amount": total,
                "description": description,
                "payment_method_id": "pix",
                "payer": {
                    "email": "pagador@example.com",
                },
            }

            result = sdk.payment().create(payment_data)
            response = result.get("response", {})

            if response.get("status") == "pending" and "point_of_interaction" in response:
                data = response["point_of_interaction"]["transaction_data"]
                link_pagamento = data.get("ticket_url", "")
    except Exception:
        # Se der qualquer erro na geração do Pix, segue sem o link
        link_pagamento = ""

    msg = (
        f"Olá {nome_resp}, tudo bem?\n\n"
        f"O seu filho(a) {nome_atl} está com as seguintes pendências financeiras:\n"
        f"- Mensalidades de: {meses_str}\n"
        f"- Valor total em aberto: R$ {total:,.2f}\n\n"
    )

    if link_pagamento:
        msg += (
            "Para facilitar, você pode pagar todas essas mensalidades no link abaixo:\n"
            f"{link_pagamento}\n\n"
            "Se preferir, fale com a coordenação para combinar outra forma de pagamento."
        )
    else:
        msg +=(
            "Fale com a coordenação para combinar a melhor forma de quitação, "
            "ou informe um comprovante após o pagamento."
        )

    telefone = (atleta.responsavel_telefone or atleta.telefone or "").strip()
    telefone_digits = "".join(ch for ch in telefone if ch.isdigit())

    # Normaliza: mantém apenas os últimos 11 dígitos (DDD + número)
    if len(telefone_digits) > 11:
        telefone_digits = telefone_digits[-11:]

    if len(telefone_digits) < 10:
        flash("Não há telefone válido cadastrado para este responsável.", "danger")
        return redirect(url_for("financeiro.resumo"))

    # Formato: wa.me/55DDDNÚMERO?text=...
    url = f"https://wa.me/55{telefone_digits}?text={quote(msg)}"
    return redirect(url)


# =========================
# PIX INDIVIDUAL (por cobrança)
# =========================

@financeiro_bp.route("/gerar-pix/<int:conta_id>")
@login_required
def gerar_pix(conta_id):
    if not _require_admin():
        return redirect(url_for("financeiro.resumo"))

    try:
        import mercadopago
    except ImportError:
        flash("Biblioteca 'mercadopago' não encontrada. Adicione ao requirements.txt.", "danger")
        return redirect(url_for("financeiro.resumo"))

    access_token = os.getenv("MERCADOPAGO_ACCESS_TOKEN")
    if not access_token:
        flash("MERCADOPAGO_ACCESS_TOKEN não configurado nas variáveis de ambiente.", "danger")
        return redirect(url_for("financeiro.resumo"))

    conta = ContaReceber.query.get_or_404(conta_id)
    if conta.status == "PAGO":
        flash("Esta cobrança já está marcada como PAGA.", "info")
        return redirect(url_for("financeiro.resumo"))

    sdk = mercadopago.SDK(access_token)

    description = conta.descricao or f"Mensalidade {conta.id}"
    value = float(conta.valor or 0)

    payment_data = {
        "transaction_amount": value,
        "description": description,
        "payment_method_id": "pix",
        "payer": {
            "email": "pagador@example.com",
        },
    }

    result = sdk.payment().create(payment_data)
    response = result.get("response", {})

    if response.get("status") != "pending" or "point_of_interaction" not in response:
        flash("Não foi possível gerar o Pix no momento.", "danger")
        return redirect(url_for("financeiro.resumo"))

    data = response["point_of_interaction"]["transaction_data"]
    qr_code = data.get("qr_code")
    qr_base64 = data.get("qr_code_base64")
    ticket_url = data.get("ticket_url")

    return render_template(
        "financeiro_pix.html",
        conta=conta,
        qr_code=qr_code,
        qr_base64=qr_base64,
        ticket_url=ticket_url,
    )


# =========================
# DESPESAS (CONTAS A PAGAR)
# =========================

@financeiro_bp.route("/despesa/nova", methods=["GET", "POST"])
@login_required
def nova_despesa():
    """
    Cadastro de despesa com:
    - parcelamento
    - data de vencimento
    - "método de pagamento" (armazenado no texto da descrição)
    """
    if not _require_admin():
        return redirect(url_for("financeiro.resumo"))

    if request.method == "POST":
        from dateutil.relativedelta import relativedelta

        fornecedor = (request.form.get("fornecedor") or "").strip()
        descricao = (request.form.get("descricao") or "").strip()
        venc_str = request.form.get("vencimento") or ""
        valor_str = request.form.get("valor") or "0"
        parcelas_str = request.form.get("parcelas") or "1"
        metodo = (request.form.get("metodo_pagamento") or "").strip()

        if not fornecedor or not venc_str:
            flash("Fornecedor e vencimento são obrigatórios.", "danger")
            return redirect(url_for("financeiro.nova_despesa"))

        try:
            vencimento = datetime.strptime(venc_str, "%Y-%m-%d").date()
        except ValueError:
            flash("Data de vencimento inválida.", "danger")
            return redirect(url_for("financeiro.nova_despesa"))

        try:
            valor = float(valor_str.replace(",", "."))
            parcelas = int(parcelas_str)
            if parcelas < 1:
                parcelas = 1
        except ValueError:
            flash("Valor ou parcelas inválidas.", "danger")
            return redirect(url_for("financeiro.nova_despesa"))

        for i in range(parcelas):
            venci_parcela = vencimento + relativedelta(months=i)
            desc_parcela = descricao or "Despesa"

            if parcelas > 1:
                desc_parcela = f"{desc_parcela} ({i+1}/{parcelas})"

            if metodo:
                desc_parcela = f"{desc_parcela} · Pagamento: {metodo}"

            conta = ContaPagar(
                fornecedor=fornecedor,
                descricao=desc_parcela,
                vencimento=venci_parcela,
                valor=valor,
                status="PENDENTE",
            )
            db.session.add(conta)

        db.session.commit()
        flash("Despesa registrada com sucesso.", "success")
        return redirect(url_for("financeiro.resumo"))

    return render_template("financeiro_despesa_form.html")


# =========================
# COBRANÇA MANUAL
# =========================

@financeiro_bp.route("/cobranca/nova", methods=["GET", "POST"])
@login_required
def nova_cobranca():
    atleta_id_param = request.args.get("atleta_id", type=int)
    atletas = Atleta.query.order_by(Atleta.nome).all()
    atleta_selecionado = None
    if atleta_id_param:
        atleta_selecionado = next((a for a in atletas if a.id == atleta_id_param), None)

    if request.method == "POST":
        from dateutil.relativedelta import relativedelta

        atleta_id = request.form.get("atleta_id")
        descricao = (request.form.get("descricao") or "").strip() or "Mensalidade"
        valor_str = request.form.get("valor") or "0"
        parcelas_str = request.form.get("parcelas") or "1"
        venc_str = request.form.get("vencimento") or ""

        if not atleta_id or not venc_str:
            flash("Selecione o atleta e a data de vencimento.", "danger")
            return redirect(url_for("financeiro.nova_cobranca"))

        try:
            atleta_id = int(atleta_id)
            valor = float(valor_str.replace(",", "."))
            parcelas = int(parcelas_str)
            if parcelas < 1:
                parcelas = 1
            vencimento = datetime.strptime(venc_str, "%Y-%m-%d").date()
        except ValueError:
            flash("Verifique valor, parcelas e vencimento.", "danger")
            return redirect(url_for("financeiro.nova_cobranca"))

        for i in range(parcelas):
            venci_parcela = vencimento + relativedelta(months=i)
            desc_parcela = (
                f"{descricao} ({i+1}/{parcelas})" if parcelas > 1 else descricao
            )
            conta = ContaReceber(
                atleta_id=atleta_id,
                descricao=desc_parcela,
                competencia=date(venci_parcela.year, venci_parcela.month, 1),
                vencimento=venci_parcela,
                valor=valor,
                status="PENDENTE",
            )
            db.session.add(conta)

        db.session.commit()
        flash("Cobrança(s) criada(s) com sucesso.", "success")
        return redirect(url_for("financeiro.resumo"))

    return render_template("financeiro_cobranca_form.html", atletas=atletas)